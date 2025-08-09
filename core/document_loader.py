from typing import List
from pathlib import Path
import PyPDF2
from docx import Document as DocxDocument
from core.models import Document

class DocumentLoader:

    @staticmethod
    def extract_flat_table_chunks(file_path: str, chunk_size: int = 3) -> List[Document]:
        doc = DocxDocument(file_path)
        all_documents = []

        tabela_id = ""
        tabela_nome = ""
        dados_formatados = []
        seen_lines = set()  # ← aqui vamos evitar duplicações

        def processar_chunk_final():
            if dados_formatados:
                for i in range(0, len(dados_formatados), chunk_size):
                    chunk_linhas = dados_formatados[i:i + chunk_size]
                    chunk_body = "\n---\n".join(chunk_linhas)
                    chunk_text = f"Tabela: {tabela_nome} ({tabela_id})\n\n{chunk_body}"
                    all_documents.append(Document(
                        id=f"{file_path}_chunk_{tabela_id}_{i}",
                        content=chunk_text,
                        metadata={"chunked": True, "tabela": tabela_nome, "codigo": tabela_id}
                    ))

        for table in doc.tables:
            rows = [
                [cell.text.strip() for cell in row.cells]
                for row in table.rows
                if any(cell.text.strip() for cell in row.cells)
            ]

            if not rows or len(rows) < 2:
                continue

            headers_detected = False

            for row in rows:
                if len(row) == 2 and row[0].startswith("NFCES"):
                    processar_chunk_final()
                    tabela_id = row[0]
                    tabela_nome = row[1]
                    dados_formatados = []
                    seen_lines = set()  # reset deduplicação para nova tabela
                    continue

                if not headers_detected and "NOME DO" in row[0].upper():
                    headers_detected = True
                    continue  # pular cabeçalho

                if headers_detected:
                    row = [cell.strip() if cell else "-" for cell in row + [""] * (10 - len(row))]

                    # pular campos irrelevantes
                    if "sem uso" in row[8].lower() or (not row[8] or row[8] == "-"):
                        continue

                    # gerar string semanticamente útil
                    nome_logico = row[0] if row[0] != "-" else None
                    nome_fisico = row[3] if row[3] != "-" else None
                    tipo_logico = row[1]
                    tipo_fisico = row[4]
                    descricao = row[8]
                    dominios = row[9] if row[9] != "-" else None

                    # montar texto base
                    linhas = []

                    if nome_logico and nome_fisico and nome_logico != nome_fisico:
                        linhas.append(f"Campo: {nome_logico} → {nome_fisico}")
                    elif nome_logico:
                        linhas.append(f"Campo: {nome_logico}")
                    elif nome_fisico:
                        linhas.append(f"Campo (físico): {nome_fisico}")

                    linhas.append(f"Descrição: {descricao}")

                    if tipo_logico != "-" or tipo_fisico != "-":
                        linhas.append(f"Tipo: {tipo_logico} / {tipo_fisico}")

                    if dominios:
                        linhas.append(f"Domínios: {dominios}")

                    texto_formatado = "\n".join(linhas)

                    # evitar duplicação exata
                    if texto_formatado not in seen_lines:
                        dados_formatados.append(texto_formatado)
                        seen_lines.add(texto_formatado)

        # Finalizar última tabela se restou algo
        processar_chunk_final()

        return all_documents

    @staticmethod
    def load_pdf(file_path: str) -> str:
        text = ""
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                text += page.extract_text() or ""
        return text

    @staticmethod
    def load_documents(paths: List[str], chunk_size: int = 4) -> List[Document]:
        docs = []
        for path in paths:
            ext = Path(path).suffix.lower()
            if ext == '.pdf':
                content = DocumentLoader.load_pdf(path)
                docs.append(Document(id=str(path), content=content))
            elif ext == '.docx':
                chunks = DocumentLoader.extract_flat_table_chunks(path, chunk_size=chunk_size)
                for i, chunk in enumerate(chunks):
                    print(f"\n--- Chunk {i} ---\n{chunk.content}")
                docs.extend(chunks)
            else:
                continue
        return docs
